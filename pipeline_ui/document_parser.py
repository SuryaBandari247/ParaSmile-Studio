"""Multi-format document parser for uploaded files.

Extracts plain text from PDF, XLSX, CSV, TXT, MD, and DOCX files
to provide supplementary context for LLM-based script conversion.
"""

from dataclasses import dataclass
from io import BytesIO
from pathlib import PurePosixPath
from typing import Optional

SUPPORTED_EXTENSIONS = {".pdf", ".xlsx", ".csv", ".txt", ".md", ".docx"}


@dataclass
class ParsedDocument:
    """Result of parsing a single uploaded file."""

    filename: str
    text: str
    char_count: int
    success: bool
    error: Optional[str] = None


def parse_file(filename: str, content: bytes) -> ParsedDocument:
    """Extract text from a single file based on its extension.

    Returns a ParsedDocument with success=False and error message
    if the format is unsupported or extraction fails.
    """
    ext = PurePosixPath(filename).suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        return ParsedDocument(
            filename=filename,
            text="",
            char_count=0,
            success=False,
            error=f"Unsupported format: {ext}",
        )

    extractors = {
        ".pdf": _extract_pdf,
        ".xlsx": _extract_xlsx,
        ".csv": _extract_csv,
        ".txt": _extract_txt,
        ".md": _extract_txt,
        ".docx": _extract_docx,
    }

    try:
        text = extractors[ext](content)
        return ParsedDocument(
            filename=filename,
            text=text,
            char_count=len(text),
            success=True,
        )
    except Exception as e:
        return ParsedDocument(
            filename=filename,
            text="",
            char_count=0,
            success=False,
            error=str(e),
        )


def parse_files(files: list[tuple[str, bytes]]) -> list[ParsedDocument]:
    """Parse multiple files, returning results for each (including failures)."""
    return [parse_file(name, data) for name, data in files]


def concatenate_extracted_text(documents: list[ParsedDocument]) -> str:
    """Concatenate successful extractions with document boundary markers.

    Format per document:
        --- Document: filename.pdf ---
        <extracted text>
        --- End Document: filename.pdf ---

    Returns empty string for an empty list.
    """
    if not documents:
        return ""

    sections: list[str] = []
    for doc in documents:
        if doc.success:
            sections.append(
                f"--- Document: {doc.filename} ---\n"
                f"{doc.text}\n"
                f"--- End Document: {doc.filename} ---"
            )
    return "\n".join(sections)


# ---------------------------------------------------------------------------
# Internal format-specific extractors
# ---------------------------------------------------------------------------


def _extract_pdf(content: bytes) -> str:
    """Extract text from all pages of a PDF using pypdf."""
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(content))
    pages_text: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages_text.append(text)
    return "\n".join(pages_text)


def _extract_xlsx(content: bytes) -> str:
    """Extract cell values from all sheets, rows as comma-separated lines."""
    from openpyxl import load_workbook

    wb = load_workbook(BytesIO(content), read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            lines.append(",".join(str(cell) if cell is not None else "" for cell in row))
    wb.close()
    return "\n".join(lines)


def _extract_csv(content: bytes) -> str:
    """Read CSV content as plain UTF-8 text."""
    return content.decode("utf-8")


def _extract_txt(content: bytes) -> str:
    """Read plain text / markdown content as UTF-8."""
    return content.decode("utf-8")


def _extract_docx(content: bytes) -> str:
    """Extract paragraph text from a DOCX document."""
    from docx import Document

    doc = Document(BytesIO(content))
    return "\n".join(p.text for p in doc.paragraphs)
