"""Unit tests for pipeline_ui.document_parser."""

from io import BytesIO

import pytest

from pipeline_ui.document_parser import (
    ParsedDocument,
    _extract_docx,
    _extract_pdf,
    _extract_xlsx,
    concatenate_extracted_text,
    parse_file,
    parse_files,
)


# ---------------------------------------------------------------------------
# parse_file — plain text formats
# ---------------------------------------------------------------------------


class TestParseFilePlainText:
    """Tests for parse_file with UTF-8 text-based formats."""

    def test_valid_txt_file(self):
        result = parse_file("hello.txt", b"Hello world")
        assert result.success is True
        assert result.text == "Hello world"
        assert result.char_count == 11
        assert result.filename == "hello.txt"
        assert result.error is None

    def test_valid_csv_file(self):
        result = parse_file("data.csv", b"a,b,c\n1,2,3")
        assert result.success is True
        assert result.text == "a,b,c\n1,2,3"
        assert result.char_count == 11

    def test_valid_md_file(self):
        result = parse_file("readme.md", b"# Title\nContent")
        assert result.success is True
        assert result.text == "# Title\nContent"
        assert result.char_count == 15

    def test_zero_byte_txt_file(self):
        result = parse_file("empty.txt", b"")
        assert result.success is True
        assert result.text == ""
        assert result.char_count == 0


# ---------------------------------------------------------------------------
# parse_file — unsupported format
# ---------------------------------------------------------------------------


class TestParseFileUnsupported:
    def test_unsupported_extension_returns_failure(self):
        result = parse_file("image.xyz", b"\x00\x01\x02")
        assert result.success is False
        assert "Unsupported" in result.error
        assert ".xyz" in result.error
        assert result.text == ""
        assert result.char_count == 0


# ---------------------------------------------------------------------------
# parse_files — batch processing
# ---------------------------------------------------------------------------


class TestParseFiles:
    def test_returns_one_result_per_input(self):
        files = [
            ("a.txt", b"alpha"),
            ("b.xyz", b"nope"),
            ("c.csv", b"1,2"),
        ]
        results = parse_files(files)
        assert len(results) == 3
        assert results[0].success is True
        assert results[1].success is False
        assert results[2].success is True


# ---------------------------------------------------------------------------
# concatenate_extracted_text
# ---------------------------------------------------------------------------


class TestConcatenateExtractedText:
    def test_multiple_docs_produce_boundary_markers(self):
        docs = [
            ParsedDocument(filename="a.txt", text="AAA", char_count=3, success=True),
            ParsedDocument(filename="b.csv", text="BBB", char_count=3, success=True),
        ]
        result = concatenate_extracted_text(docs)
        assert "--- Document: a.txt ---" in result
        assert "--- End Document: a.txt ---" in result
        assert "--- Document: b.csv ---" in result
        assert "--- End Document: b.csv ---" in result
        # Order preserved
        assert result.index("a.txt") < result.index("b.csv")

    def test_empty_list_returns_empty_string(self):
        assert concatenate_extracted_text([]) == ""

    def test_failed_docs_are_excluded(self):
        docs = [
            ParsedDocument(filename="ok.txt", text="OK", char_count=2, success=True),
            ParsedDocument(filename="bad.xyz", text="", char_count=0, success=False, error="Unsupported"),
        ]
        result = concatenate_extracted_text(docs)
        assert "ok.txt" in result
        assert "bad.xyz" not in result


# ---------------------------------------------------------------------------
# _extract_pdf — in-memory PDF
# ---------------------------------------------------------------------------


class TestExtractPdf:
    def test_extract_pdf_blank_page(self):
        """A blank page has no text — extraction should return empty string."""
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=72, height=72)
        buf = BytesIO()
        writer.write(buf)
        pdf_bytes = buf.getvalue()

        result = _extract_pdf(pdf_bytes)
        assert isinstance(result, str)
        # Blank page yields empty text
        assert result.strip() == ""

    def test_parse_file_with_pdf(self):
        """End-to-end: parse_file dispatches to _extract_pdf."""
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=72, height=72)
        buf = BytesIO()
        writer.write(buf)

        result = parse_file("report.pdf", buf.getvalue())
        assert result.success is True


# ---------------------------------------------------------------------------
# _extract_xlsx — in-memory workbook
# ---------------------------------------------------------------------------


class TestExtractXlsx:
    def test_extract_xlsx_basic(self):
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.append(["Name", "Value"])
        ws.append(["A", "1"])
        buf = BytesIO()
        wb.save(buf)

        result = _extract_xlsx(buf.getvalue())
        assert "Name,Value" in result
        assert "A,1" in result

    def test_parse_file_with_xlsx(self):
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.append(["X", "Y"])
        buf = BytesIO()
        wb.save(buf)

        result = parse_file("data.xlsx", buf.getvalue())
        assert result.success is True
        assert "X,Y" in result.text


# ---------------------------------------------------------------------------
# _extract_docx — in-memory document
# ---------------------------------------------------------------------------


class TestExtractDocx:
    def test_extract_docx_basic(self):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Hello from docx")
        buf = BytesIO()
        doc.save(buf)

        result = _extract_docx(buf.getvalue())
        assert "Hello from docx" in result

    def test_parse_file_with_docx(self):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Paragraph one")
        doc.add_paragraph("Paragraph two")
        buf = BytesIO()
        doc.save(buf)

        result = parse_file("notes.docx", buf.getvalue())
        assert result.success is True
        assert "Paragraph one" in result.text
        assert "Paragraph two" in result.text
